from __future__ import annotations

import json
import sqlite3
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "REPORT"
OUT_PATH = REPORT_DIR / "AI_Scientist智能体系统_章稿初稿.docx"
DB_PATH = ROOT / "文献检索_agent" / "data" / "literature.sqlite"
INDEX_STATUS_PATH = ROOT / "文献分析_agent" / "data" / "index" / "index_status.json"
ROUTE_PATH = ROOT / "04_提出路线" / "data" / "route_candidates.json"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F4F6F9"
BORDER = "B8C7D9"


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    stats = collect_stats()
    index_status = read_json(INDEX_STATUS_PATH, {})
    route_summary = collect_route_summary()

    doc = Document()
    configure_document(doc)
    add_running_header_footer(doc)
    add_title(doc)
    add_intro_note(doc, stats, index_status)

    section_41(doc)
    section_42(doc)
    section_43(doc, stats)
    section_44(doc, stats, index_status)
    section_45(doc, route_summary)
    section_46(doc, stats, index_status, route_summary)
    section_47(doc)

    audit_docx_structure(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


def collect_stats() -> dict[str, int | list[tuple[str, int]] | list[tuple[str, str, str, str]]]:
    stats: dict[str, int | list[tuple[str, int]] | list[tuple[str, str, str, str]]] = {}
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        count_queries = {
            "papers": "SELECT COUNT(*) FROM papers",
            "source_records": "SELECT COUNT(*) FROM source_records",
            "search_runs": "SELECT COUNT(*) FROM search_runs",
            "goals": "SELECT COUNT(*) FROM scientific_goals",
            "rounds": "SELECT COUNT(*) FROM exploration_rounds",
            "round_candidates": "SELECT COUNT(*) FROM round_candidates",
            "pdf_assets": "SELECT COUNT(*) FROM pdf_assets",
            "paper_conversions": "SELECT COUNT(*) FROM paper_conversions",
            "paper_cards": "SELECT COUNT(*) FROM paper_cards",
            "knowledge_nodes": "SELECT COUNT(*) FROM knowledge_nodes",
            "knowledge_edges": "SELECT COUNT(*) FROM knowledge_edges",
            "wiki_pages": "SELECT COUNT(*) FROM wiki_pages",
        }
        for key, query in count_queries.items():
            stats[key] = int(conn.execute(query).fetchone()[0])
        stats["access_status"] = [
            (str(row[0]), int(row[1]))
            for row in conn.execute(
                "SELECT access_status, COUNT(*) FROM access_records "
                "GROUP BY access_status ORDER BY COUNT(*) DESC"
            )
        ]
        stats["round_status"] = [
            (str(row[0]), int(row[1]))
            for row in conn.execute(
                "SELECT status, COUNT(*) FROM exploration_rounds "
                "GROUP BY status ORDER BY COUNT(*) DESC"
            )
        ]
        stats["goals_list"] = [
            (
                str(row["id"]),
                str(row["title"] or ""),
                str(row["description"] or ""),
                str(row["default_target_count"] or ""),
            )
            for row in conn.execute(
                "SELECT id, title, description, default_target_count "
                "FROM scientific_goals ORDER BY updated_at DESC, id DESC"
            )
        ]
    return stats


def read_json(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return default


def collect_route_summary() -> dict[str, object]:
    data = read_json(ROUTE_PATH, {"runs": []})
    runs = data.get("runs") if isinstance(data, dict) else []
    latest = runs[0] if runs else {}
    routes = latest.get("routes", []) if isinstance(latest, dict) else []
    return {
        "saved_runs": len(runs),
        "latest_question": latest.get("question_title", "") if isinstance(latest, dict) else "",
        "latest_mode": (latest.get("metadata") or {}).get("mode", "") if isinstance(latest, dict) else "",
        "latest_error": (latest.get("metadata") or {}).get("error", "") if isinstance(latest, dict) else "",
        "routes": routes[:4] if isinstance(routes, list) else [],
    }


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.85)
    section.right_margin = Cm(2.85)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after, bold in [
        ("Heading 1", 16, BLUE, 18, 10, True),
        ("Heading 2", 13, BLUE, 12, 6, True),
        ("Heading 3", 12, DARK_BLUE, 8, 4, True),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(10)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color=None, east_asia="宋体") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("基于多智能体的跨学科文献知识挖掘与研究路线生成系统")
    set_run_font(run, size=9, color=MUTED, east_asia="宋体")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    add_field(footer_p, "PAGE")


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("第四章 基于多智能体的跨学科文献知识挖掘与研究路线生成系统")
    set_run_font(run, size=18, bold=True, color=RGBColor(0, 0, 0), east_asia="黑体")


def add_intro_note(doc: Document, stats: dict, index_status: dict) -> None:
    rows = [
        ("章节定位", "毕业论文单章初稿，面向“智能体系统开发与应用验证”章节，可后续并入完整论文模板。"),
        ("当前数据底座", f"检索库已收录论文 {stats['papers']} 篇、来源记录 {stats['source_records']} 条、PDF 资产 {stats['pdf_assets']} 个。"),
        ("知识组织状态", f"已形成 Paper Card {stats['paper_cards']} 个、知识节点 {stats['knowledge_nodes']} 个、知识边 {stats['knowledge_edges']} 条、Wiki 条目 {stats['wiki_pages']} 个。"),
        ("RAG 索引状态", f"当前索引文档 {index_status.get('documents', 0)} 篇、chunks {index_status.get('chunks', 0)} 个、embedding {index_status.get('embeddings', 0)} 条，模式为 {index_status.get('embedding_mode', '未记录')}。"),
    ]
    add_table(
        doc,
        "表4.1 本章写作所依据的当前项目状态",
        ["项目", "内容"],
        rows,
        widths=[1.35, 5.15],
    )


def h1(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def h3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    for run in p.runs:
        set_run_font(run, size=11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run_font(run, size=11)


def section_41(doc: Document) -> None:
    h1(doc, "4.1 研究背景")
    para(
        doc,
        "不同学科的研究问题都可能涉及理论机制、实验条件、数据方法和结果指标之间的多尺度关联。"
        "对于“某个关键变量如何影响目标结果”这类科学问题，单一文献往往只能给出局部证据，真正可用于研究决策的答案需要跨论文比较"
        "研究对象、计算模型、实验条件与评价指标。"
    )
    para(
        doc,
        "传统文献调研主要依赖研究者手工检索、阅读和摘录。该方式在小范围主题上具有较高可信度，但面对持续增长的论文数量、"
        "异构数据库来源以及 PDF 正文、图表、补充信息混杂的证据形态时，容易出现覆盖不足、证据链断裂和更新滞后的问题。"
        "此外，研究中的关键知识并不总是以结构化表格出现，许多重要信息存在于机理讨论、模型假设、边界条件和负面结论之中。"
        "因此，构建一个能够将文献获取、证据组织、科学问题归纳和研究路线生成连接起来的智能体系统，是提高科研效率的重要方向。"
    )
    para(
        doc,
        "本章围绕当前搭建的 AI Scientist 原型系统，介绍一个面向跨学科研究的多智能体工作流。该系统并不试图用大语言模型直接替代研究者判断，"
        "而是将确定性的数据入库、去重、状态流转、RAG 索引和报表统计作为基础，将大语言模型限定在查询扩展、边界复核、候选重排、"
        "问题归纳和路线草拟等辅助环节中。通过这种设计，系统既能够利用语言模型的语义组织能力，又保留对文献来源、PDF 可得性、"
        "证据编号和运行状态的可追溯控制。"
    )
    para(
        doc,
        "本章的写作重点是系统开发与初步应用验证，而非某一具体学科的完整机制研究。换言之，本章关注的是如何把分散文献转化为可检索、"
        "可追踪、可讨论并能继续进入路线设计的研究上下文；至于具体研究对象与机制的最终结论，仍需要后续更大规模语料、"
        "更精细的计算验证和实验反馈来支撑。"
    )


def section_42(doc: Document) -> None:
    h1(doc, "4.2 多智能体系统总体架构")
    para(
        doc,
        "当前系统采用分阶段、多 Agent 协同的方式组织 AI Scientist 工作流。整体链路从文献检索开始，经过 PDF 解析和知识组织，"
        "再进入科学问题归纳，最后生成研究路线候选。四个阶段之间通过 SQLite 数据库、本地 PDF 文件库、Markdown/JSON 解析产物、"
        "RAG 索引、Paper Cards、知识图谱和 Wiki 条目传递上下文。文献检索 Agent 提供统一的 8765 本地网页入口，后续模块通过桥接层接入该前端，"
        "使研究者能够在同一个交互界面中完成问题创建、轮次规划、PDF 获取、文献分析、问题归纳和路线查看。"
    )
    add_table(
        doc,
        "表4.2 系统四阶段 Agent 分工",
        ["阶段", "核心职责", "主要输入", "主要输出"],
        [
            (
                "文献检索 Agent",
                "多源召回、元数据规范化、去重、开放获取状态解析、PDF 下载、人工确认轮次和质量审计。",
                "科学问题、领域查询配置、外部文献源。",
                "SQLite 文献库、PDF 文件库、轮次候选、报表和仪表盘。",
            ),
            (
                "文献分析 Agent",
                "PDF 转 Markdown/JSON，构建 RAG 索引、Paper Cards、轻量知识图谱和 Wiki。",
                "检索数据库、PDF 文件库、轮次 PDF。",
                "解析文档、FTS/向量索引、证据化卡片、图谱与 Wiki 条目。",
            ),
            (
                "科学问题归纳 Agent",
                "结合检索问题、证据缺口、Paper Cards 和 Wiki，对话式收敛可验证科学问题。",
                "检索库、分析产物、研究者输入。",
                "对话数据库、问题候选、后续追问建议。",
            ),
            (
                "研究路线候选 Agent",
                "把科学问题与证据上下文转化为路线候选，给出材料/结构、变量、验证方式、风险和下一步。",
                "科学问题、Paper Cards、Wiki、证据缺口。",
                "候选路线 JSON 和前端可编辑草稿。",
            ),
        ],
        widths=[1.15, 2.35, 1.45, 1.55],
    )
    para(
        doc,
        "系统架构的一个关键原则是将“事实状态”和“语言推理”分离。论文记录、DOI、PDF 路径、下载状态、轮次状态、Paper Card 和 Wiki 等对象"
        "均以结构化文件或数据库表保存；LLM 的输出则被视为可审阅的候选判断，而不是直接覆盖事实表的唯一依据。例如，在文献选择和路线生成阶段，"
        "外部 LLM 可以参与候选重排或 JSON 草拟，但当模型未配置、调用失败或返回不合法结果时，系统会回退到本地规则生成草稿，"
        "并在元数据中保留失败原因。"
    )
    h2(doc, "4.2.1 数据流与模块边界")
    para(
        doc,
        "第一阶段的检索数据库是整个系统的主索引，记录 papers、source_records、access_records、round_candidates、manual_download_tasks、"
        "paper_conversions 等表。第二阶段的分析模块读取该数据库和 PDF 文件库，将可解析论文转换为 Markdown/JSON，并将转换状态写回检索数据库，"
        "从而让前端能够统一展示 PDF 是否已完成分析。第三阶段不直接重新解析 PDF，而是读取分析模块生成的 Paper Cards、Wiki 条目和知识图谱，"
        "把它们压缩成适合对话推理的上下文。第四阶段进一步读取科学问题与证据上下文，生成可供研究者修改和执行的路线草稿。"
    )
    h2(doc, "4.2.2 人机协同与可追溯性设计")
    para(
        doc,
        "本系统没有把 AI Scientist 设计为完全自动闭环，而是保留了多个研究者确认节点。文献轮次需要人工批准后才进入 PDF 获取；"
        "非开放或下载失败论文被转化为手动下载任务；科学问题归纳通过对话逐步细化；研究路线候选以 JSON 草稿保存，允许研究者继续编辑。"
        "这种设计适合早期科研探索场景：系统负责提高覆盖率、组织证据和暴露缺口，研究者负责判断研究价值、实验可行性和最终取舍。"
    )


def section_43(doc: Document, stats: dict) -> None:
    h1(doc, "4.3 文献检索与语料构建 Agent")
    para(
        doc,
        "文献检索 Agent 是系统的入口模块，目标是围绕给定科学问题构建可复现的文献语料库。"
        "该模块采用单进程多 Agent 的 CLI 与网页原型，内部由领域查询规划、数据源发现、元数据清洗、去重、相关性判断、"
        "开放获取解析、PDF 下载、质量审计、报表导出和 HTML 仪表盘等子 Agent 组成。其设计重点不是一次性得到最终答案，"
        "而是支持按轮次、小批量、可回溯地扩展文献集合。"
    )
    h2(doc, "4.3.1 多源召回与查询规划")
    para(
        doc,
        "在查询规划阶段，系统根据科学问题和领域配置生成一组检索 query，并可选择 smoke、pilot 和 full 等不同规模。"
        "查询词由研究者给出的研究对象、关键变量、方法和目标指标生成。SourceDiscoveryAgent 负责从外部数据源召回候选记录，"
        "并把数据源失败、限流、超时和 HTTP 错误分类写入 source_failures，便于后续审计覆盖率。"
    )
    para(
        doc,
        "召回后，MetadataNormalizeAgent 将标题、作者、年份、DOI、来源链接等字段规范化，DeduplicationAgent 再根据 DOI、标题相似度和来源线索合并重复记录。"
        "相关性判断采用确定性规则作为主线，LLMRelevanceReviewAgent 只复核边界样本，并记录复核结论或跳过原因。"
        "这种安排降低了大模型幻觉对核心语料库的影响，也使每篇论文为何进入或未进入候选池更容易解释。"
    )
    h2(doc, "4.3.2 开放获取、PDF 获取与人工确认轮次")
    para(
        doc,
        "OAResolverAgent 负责补全 DOI URL、出版商页面、开放获取状态和 PDF URL；PdfDownloadAgent 只下载开放获取或预印本文献，"
        "对于闭源论文不绕过权限，而是保留 DOI、出版商页面和来源线索。RoundPlanningAgent 在每轮检索后生成候选集合，"
        "LiteratureSelectionAgent 根据目标问题、材料标签、证据等级和多样性策略选择本轮文献。系统支持外部 LLM 在真实候选池内重排，"
        "但模型只能选择已有 paper_id，不能生成不存在的论文或 DOI；当模型不可用时，则由规则排序补足。"
    )
    para(
        doc,
        "当前数据库中共有 "
        f"{stats['goals']} 个科学问题和 {stats['rounds']} 个探索轮次，轮次状态均记录为 awaiting_manual_pdfs。"
        "这说明系统已经完成候选规划、用户批准和自动获取阶段，并进入需要人工补充闭源 PDF 的状态。"
        "这种状态流转虽然还不是全自动完成，但符合科研文献获取的实际约束：开放文献自动处理，受版权限制的文献交由研究者合法获取。"
    )
    access_rows = [(status_label(k), str(v)) for k, v in stats["access_status"]]
    add_table(
        doc,
        "表4.3 当前 access_records 中的文献可得性状态",
        ["访问状态", "记录数"],
        access_rows,
        widths=[4.5, 2.0],
    )
    h2(doc, "4.3.3 质量审计与报表输出")
    para(
        doc,
        "检索模块不仅保存论文数据，也输出方法指标和质量审计结果。报表中包含 DOI 覆盖率、OA 覆盖率、PDF URL 覆盖率、下载 PDF 覆盖率、"
        "去重压缩率、数据源失败统计和 LLM 边界复核记录等。当前项目目录中保留了一次小规模报表样例，其中记录 papers 为 58、source_records 为 71、"
        "downloaded 为 10、source_failures 为 22、llm_reviews 为 62、audit_findings 为 15。该历史报表与当前 SQLite 总库统计不同，"
        "反映的是某一次 report 命令导出的运行快照；本章后续系统验证以当前 SQLite 库统计为准。"
    )


def section_44(doc: Document, stats: dict, index_status: dict) -> None:
    h1(doc, "4.4 文献解析与证据化知识组织 Agent")
    para(
        doc,
        "文献分析 Agent 是系统的第二阶段，负责把检索模块获得的 PDF 转化为后续推理可用的证据对象。"
        "该模块默认读取检索数据库、PDF 文件库和轮次 PDF 目录，输出 parsed_papers、index、cards、graph 和 wiki 等数据目录。"
        "当前版本使用 PyMuPDF 进行本地 PDF 文本抽取，并通过规则识别标题、摘要、章节、正文段落、图注、表格标题和参考文献。"
        "对于扫描版或文本不可抽取 PDF，系统会记录失败或低质量结果；OCR 与图表多模态解析仍属于后续扩展方向。"
    )
    h2(doc, "4.4.1 PDF 转换与 RAG 索引")
    para(
        doc,
        "PDF 转换阶段将论文内容保存为 Markdown 和 JSON 两类文件。Markdown 便于研究者直接阅读，JSON 则保留结构化字段和 evidence id。"
        "随后，CorpusIndexAgent 从 parsed_papers 中构建 SQLite FTS5 全文索引，并写入本地哈希向量；当外部 embedding 服务未配置时，"
        "系统自动使用 local_hash_fallback 作为兜底。当前索引状态显示，系统已经索引 "
        f"{index_status.get('documents', 0)} 篇文档，生成 {index_status.get('chunks', 0)} 个文本块和 {index_status.get('embeddings', 0)} 条 embedding。"
    )
    para(
        doc,
        "从研究用途看，RAG 层承担的是“可回溯检索”而不是最终结论生成。科学问题归纳或路线生成时，系统可以先基于关键词或语义相似度找到相关文本块，"
        "再将 evidence id、主题标签、方法标签和结果指标交给上层 Agent 组织回答。这样可以减少语言模型脱离原文证据自由发挥的风险。"
    )
    h2(doc, "4.4.2 Paper Cards、知识图谱与 Wiki")
    para(
        doc,
        "PaperCardAgent 将单篇论文压缩为研究对象、主题实体、方法、结果指标、结论、局限和 evidence ids 等字段。"
        "主题标签由论文标题、摘要和正文自动抽取，方法标签和可量化结果则用于建立可追溯的证据关联。"
    )
    para(
        doc,
        "KnowledgeGraphAgent 在 Paper Cards 基础上生成轻量知识图谱，节点类型包括 wiki_topic、paper、claim、material、method、property 和 evidence。"
        "当前系统已生成 "
        f"{stats['knowledge_nodes']} 个节点和 {stats['knowledge_edges']} 条边。WikiAgent 再围绕 topic 生成 Markdown/JSON 条目，"
        f"当前已有 {stats['wiki_pages']} 个 Wiki 条目，覆盖当前语料中抽取出的研究实体、方法和结果指标。"
    )
    add_table(
        doc,
        "表4.4 文献分析阶段的证据化产物",
        ["产物", "当前数量", "功能"],
        [
            ("Paper Cards", str(stats["paper_cards"]), "把单篇论文压缩为材料、方法、性能、结论和 evidence ids。"),
            ("知识节点", str(stats["knowledge_nodes"]), "表达 topic、paper、claim、material、method、property、evidence 等对象。"),
            ("知识边", str(stats["knowledge_edges"]), "连接论文、证据、材料、方法和性能指标之间的关系。"),
            ("Wiki 条目", str(stats["wiki_pages"]), "围绕主题生成可读条目，并标出证据和待补充问题。"),
        ],
        widths=[1.45, 1.0, 4.05],
    )
    para(
        doc,
        "需要强调的是，当前分析模块仍处于样例贯通阶段。虽然数据库表和文件结构已经支持从 PDF 到 RAG、Paper Card、图谱和 Wiki 的端到端流程，"
        "但已完成深度解析的论文数量还较少，无法代表任何研究领域的完整知识覆盖。后续工作应扩大 PDF 解析规模，增加 OCR 与图表解析，"
        "并引入人工抽样复核来评估 Paper Cards 和 Wiki 条目的准确率。"
    )


def section_45(doc: Document, route_summary: dict) -> None:
    h1(doc, "4.5 科学问题归纳与研究路线生成 Agent")
    para(
        doc,
        "第三阶段的科学问题归纳 Agent 面向研究者对话场景。它读取文献检索阶段创建的科学问题、下一轮 query、证据缺口，以及文献分析阶段生成的 Paper Cards、"
        "Wiki 条目和知识图谱，帮助研究者把宽泛问题压缩为可研究、可验证、可继续进入路线设计的科学问题。系统提示明确要求 Agent 不编造文献、DOI 或实验结果，"
        "当证据不足时必须说明“证据不足”，并优先引用上下文中的 evidence ids、topic、材料体系、方法和性能指标。"
    )
    para(
        doc,
        "在当前样例中，研究者提出的问题是“关键变量如何通过可验证机制影响目标结果”。该问题本身包含研究对象、局部条件、计算或测量方法、"
        "以及可验证指标等多个未展开维度。科学问题归纳 Agent 的作用不是直接给出单一答案，而是引导研究者明确研究边界、目标结果和验证手段，"
        "例如将问题改写为“在给定研究对象中，哪些局部条件或干预因素能够提升目标结果，并如何通过计算、实验或文献证据验证”。"
    )
    h2(doc, "4.5.1 对话式问题收敛")
    para(
        doc,
        "科学问题归纳模块维护独立的 question_synthesis.sqlite 对话数据库，保存 session 和 messages。模块初始化时会先把检索阶段问题和证据缺口转化为一条"
        "“文献检索 Agent”消息，再由科学问题归纳 LLM 或本地规则生成初始回答。后续研究者输入会与最近消息和当前证据上下文一起传入模型，"
        "返回答案及“后续可追问”建议。这样，问题收敛过程被保存为可回看记录，而不是一次性聊天结果。"
    )
    h2(doc, "4.5.2 路线候选生成")
    para(
        doc,
        "第四阶段的研究路线候选 Agent 读取科学问题、证据缺口、候选论文、Paper Cards 和 Wiki 条目，输出严格 JSON 格式的 routes 数组。"
        "每条路线包含核心思路、候选材料或结构、关键变量、验证方式、证据依据、主要风险、优先级理由和下一步行动。"
        "当外部 LLM 已配置且调用成功时，系统优先使用模型生成结构化路线；当模型未配置或调用失败时，系统使用本地规则生成可编辑草稿，并在 metadata 中记录模式。"
    )
    routes = route_summary.get("routes", [])
    if routes:
        rows = []
        for item in routes[:4]:
            rows.append(
                (
                    str(item.get("rank", "")),
                    str(item.get("title", "")),
                    ", ".join(str(v) for v in item.get("candidates", [])[:4]),
                    str(item.get("priority", "")),
                )
            )
        add_table(
            doc,
            "表4.5 当前样例问题生成的研究路线候选",
            ["排序", "路线标题", "候选材料/结构", "优先级"],
            rows,
            widths=[0.65, 2.45, 2.5, 0.9],
        )
    para(
        doc,
        "当前路线候选文件中保存了 "
        f"{route_summary.get('saved_runs', 0)} 次路线生成记录，最近一次问题为“{route_summary.get('latest_question', '')}”。"
        f"最近一次 metadata mode 为 {route_summary.get('latest_mode', '未记录')}。"
        "在该样例中，外部 LLM 请求因证书校验问题失败，系统回退到本地规则草稿。"
        "这不是完整智能推理能力的最终验证，而是说明系统已经具备失败可记录、结果可兜底、草稿可继续编辑的工程韧性。"
    )


def section_46(doc: Document, stats: dict, index_status: dict, route_summary: dict) -> None:
    h1(doc, "4.6 系统初步验证与运行结果")
    para(
        doc,
        "为了验证该多智能体系统是否形成端到端链路，本研究对当前项目目录中的数据库、分析产物和路线候选文件进行了统计。"
        "验证目标不是证明系统已经完成大规模知识发现，而是检查四个阶段是否能够围绕同一科学问题传递上下文，并生成可追溯的中间产物。"
    )
    add_table(
        doc,
        "表4.6 当前 SQLite 与分析产物统计",
        ["类别", "数量", "说明"],
        [
            ("论文记录", str(stats["papers"]), "papers 表中的去重论文条目。"),
            ("来源记录", str(stats["source_records"]), "source_records 表中的多源召回记录。"),
            ("检索运行", str(stats["search_runs"]), "search_runs 表中的检索执行记录。"),
            ("科学问题", str(stats["goals"]), "scientific_goals 表中保存的研究问题。"),
            ("探索轮次", str(stats["rounds"]), "exploration_rounds 表中保存的轮次。"),
            ("轮次候选", str(stats["round_candidates"]), "round_candidates 表中的候选论文。"),
            ("PDF 资产", str(stats["pdf_assets"]), "pdf_assets 表中记录的本地 PDF。"),
            ("转换记录", str(stats["paper_conversions"]), "paper_conversions 表中的 PDF 转换状态。"),
            ("Paper Cards", str(stats["paper_cards"]), "paper_cards 表或分析产物中的卡片记录。"),
            ("知识节点", str(stats["knowledge_nodes"]), "knowledge_nodes 表中的图谱节点。"),
            ("知识边", str(stats["knowledge_edges"]), "knowledge_edges 表中的图谱关系。"),
            ("Wiki 条目", str(stats["wiki_pages"]), "wiki_pages 表或 wiki 目录中的主题条目。"),
        ],
        widths=[1.45, 0.85, 4.2],
    )
    para(
        doc,
        "从统计结果看，系统已具备三个层面的初步能力。第一，语料构建层已经形成包含数百条来源记录和两百余篇去重论文的文献库，"
        "并能追踪开放获取、预印本、闭源 DOI 和本地 PDF 资产。第二，知识组织层已经跑通从 PDF 到文本块、embedding、Paper Card、知识图谱和 Wiki 的流程。"
        "第三，研究交互层已经能够围绕样例科学问题生成对话上下文和路线候选，并把结果保存为可复用文件。"
    )
    h2(doc, "4.6.1 样例问题运行结果")
    para(
        doc,
        "当前样例问题聚焦“关键变量如何通过可验证机制影响目标结果”。路线候选模块基于已有 Paper Card 与 Wiki 证据，生成机制优先、"
        "候选体系、条件优化和数据驱动等路线。其中，机制优先路线强调先解释变量到结果的因果链，候选体系路线强调在可行研究对象中统一指标排序，"
        "条件优化路线强调可控变量和对照设计，数据驱动路线则强调构建小样本可解释排序器。"
    )
    para(
        doc,
        "这些路线并不构成学科结论本身，而是把文献证据和研究偏好转化为可执行计划。其价值在于：研究者可以沿着路线进一步选择研究对象、"
        "设计计算、实验或数据分析方案，或将缺失证据转化为下一轮检索 query。"
    )
    h2(doc, "4.6.2 当前局限")
    add_bullets(
        doc,
        [
            "已深度解析的 PDF 数量仍然较少，Paper Card 和 Wiki 的覆盖范围不足以支持领域级结论。",
            "当前 embedding 采用 local_hash_fallback，适合作为本地兜底，但语义检索质量仍需外部或训练型向量模型评估。",
            "扫描版 PDF、图表、补充信息和材料结构图尚未纳入多模态解析，可能遗漏关键性能和结构信息。",
            "路线生成阶段已经具备 LLM 调用接口，但当前样例因 SSL 证书校验失败进入规则兜底，尚不能作为外部 LLM 路线质量的验证。",
            "系统尚未形成与第一性原理计算、微磁模拟或实验平台的自动闭环，后续路线仍需人工转化为具体计算或实验任务。",
        ],
    )
    para(
        doc,
        "因此，本章应将当前系统定位为“端到端原型和样例贯通”，而非成熟的自动科学发现系统。"
        "这种如实定位有助于后续工作围绕覆盖率、准确率、可追溯性和闭环验证逐步增强，而不是过早夸大单次样例输出的科研结论。"
    )


def section_47(doc: Document) -> None:
    h1(doc, "4.7 本章小结")
    para(
        doc,
        "本章围绕科研活动中文献证据分散、跨尺度知识组织困难和人工归纳效率不足的问题，介绍了一个基于多智能体的 AI Scientist 原型系统。"
        "该系统以文献检索 Agent 为入口，完成多源召回、元数据规范化、去重、开放获取状态解析、PDF 获取和质量审计；以文献分析 Agent 为知识组织层，"
        "将 PDF 解析为 Markdown/JSON、RAG 索引、Paper Cards、轻量知识图谱和 Wiki；以科学问题归纳 Agent 和研究路线候选 Agent 为交互层，"
        "将检索问题、证据缺口和文献证据转化为可讨论的科学问题和可执行路线草稿。"
    )
    para(
        doc,
        "当前项目已经形成从文献获取到问题收敛、路线生成的初步闭环。数据库统计表明，系统已保存两百余篇论文记录、数百条来源记录、"
        "多个探索轮次和数十个 PDF 资产；分析模块已经跑通知识卡片、知识图谱和 Wiki 的生成流程；路线模块也能够在 LLM 调用失败时保留错误元数据并生成本地规则草稿。"
        "这些结果说明，多智能体方式适合将科研文献工作拆分为可审计、可回溯、可人工介入的连续流程。"
    )
    para(
        doc,
        "后续工作主要包括四个方面：第一，扩大 PDF 解析规模并补充 OCR、图表和补充信息解析能力；第二，引入人工抽样复核和证据一致性评分，"
        "系统评估 Paper Card、Wiki 和路线候选的准确率；第三，改进 RAG 向量表示与跨文献检索质量，使科学问题归纳能够使用更完整的证据上下文；"
        "第四，将路线候选进一步连接第一性原理计算、微磁模拟或实验设计工具，形成从文献证据到可验证任务的闭环。"
        "在这些能力完善后，该系统有望成为跨学科研究中面向文献知识挖掘、科学问题收敛和研究路线设计的可持续智能基础设施。"
    )


def status_label(status: str) -> str:
    mapping = {
        "closed_access_has_doi": "闭源但有 DOI",
        "preprint_pdf": "预印本 PDF",
        "oa_pdf_available": "开放获取 PDF 可用",
        "oa_no_pdf_url": "开放获取但未发现 PDF URL",
        "downloaded_oa_pdf": "已下载开放获取 PDF",
        "closed_access_missing_doi": "闭源且缺失 DOI",
    }
    return mapping.get(status, status)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[tuple], widths: list[float]) -> None:
    p = doc.add_paragraph(caption, style="Caption")
    for run in p.runs:
        set_run_font(run, size=10, color=MUTED)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        shade_cell(hdr[idx], LIGHT_FILL)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr[idx].text = ""
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=10.5, bold=True, color=DARK_BLUE, east_asia="黑体")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if len(str(value)) < 8 and idx != len(row) - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
    set_cell_margins(table)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.first_line_indent = Pt(0)


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(9360))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def audit_docx_structure(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    # Keep the generated file as a single chapter, not a cover/document packet.
    if len(doc.sections) != 1:
        raise RuntimeError("Unexpected extra sections in chapter draft.")


if __name__ == "__main__":
    main()
